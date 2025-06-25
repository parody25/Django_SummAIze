import os
import json
from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
import fitz
from docx import Document
from openai import OpenAI
from pydantic import BaseModel
import pandas as pd
import markdown
from bs4 import BeautifulSoup

from llama_index.llms.openai import OpenAI as LlamaOpenAI
from llama_index.embeddings.openai import OpenAIEmbedding as LlamaOpenAIEmbedding
from llama_index.core import Settings
from llama_index.core import (
    load_index_from_storage,
    StorageContext,
)

from .utils import load_keywords, get_keywords_for_pdf
from dotenv import load_dotenv

from .data_models.credit_application import CreditApplication


load_dotenv()

open_API=os.getenv('OPENAI_API_KEY')
llama_cloud_api = os.getenv('LLAMA_CLOUD_API')

embed_model = LlamaOpenAIEmbedding(model="text-embedding-3-small")
llm = LlamaOpenAI(model="gpt-4o", api_key=open_API)
Settings.llm = llm
Settings.embed_model = embed_model

#Embedding and retrieval handling

def load_embeddings(application_id, text_embeddings):
    embeddings_file = f"embeddings\\{application_id}_embeddings.pkl"
    #if os.path.exists(embeddings_file):
    #    return FAISS.load_local(embeddings_file, embeddings=text_embeddings, allow_dangerous_deserialization=True)
    if os.path.exists(embeddings_file):
        print(f"Loading existing Faiss index: {embeddings_file}")
        #vector_store = FaissVectorStore.from_persist_dir(embeddings_file)
        #storage_context = StorageContext.from_defaults(vector_store=vector_store, persist_dir=embeddings_file)
        storage_context = StorageContext.from_defaults(persist_dir=embeddings_file)
        vector_db = load_index_from_storage(storage_context=storage_context)
        return vector_db


                      
    return None

def save_embeddings(application_id, embeddings, text_embeddings=OpenAIEmbeddings()):
    embeddings_file = f"embeddings\\{application_id}_embeddings.pkl"
    embeddings.save_local(embeddings_file)

    if os.path.exists(f"{embeddings_file}.faiss") and os.path.exists(f"{embeddings_file}.pkl"):
        return FAISS.load_local(embeddings_file, embeddings=text_embeddings, allow_dangerous_deserialization=True)
    return None

def validate_pdf(files):
    valid_pdfs = []
    invalid_pdfs = []
    keywords_data = load_keywords()
    keyword_dict = keywords_data["keywords"]
    validation_threshold = keywords_data["validation_threshold"]

    for file in files:
        filename = os.path.basename(file.name)
        ext = os.path.splitext(filename)[-1].lower()
        text_content = ""

        # Extract text based on file type
        if ext == ".pdf":
            pdf_reader = PdfReader(file)
            text_content = " ".join([page.extract_text() or "" for page in pdf_reader.pages])
        elif ext == ".docx":
            doc = Document(file)
            text_content = "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".xlsx":
            df = pd.read_excel(file, sheet_name=None)
            text_content = "\n".join(["\n".join(sheet_df.astype(str).values.flatten()) for sheet_df in df.values()])
        else:
            print(f"Invalid file format: {filename}")
            invalid_pdfs.append(file)
            continue

        # Validate using keywords
        keywords = get_keywords_for_pdf(filename, keyword_dict)
        matching_keywords = [keyword.lower() for keyword in keywords if keyword.lower() in text_content.lower()]
        
        if keywords and matching_keywords:
            percentage_match = (len(matching_keywords) / len(keywords)) * 100
        else:
            percentage_match = 0

        if percentage_match >= validation_threshold:
            valid_pdfs.append((file, keywords))
        else:
            invalid_pdfs.append(file)

    return valid_pdfs, invalid_pdfs

def process_question(question, vectorstore, llm):
    #docs = vectorstore.similarity_search(query=question, k=3)
    #docs = vectorstore.max_marginal_relevance_search(question, k=5, fetch_k=10)
    #chain = load_qa_chain(llm=llm, chain_type="stuff")
    #with get_openai_callback() as cb:
        #response = chain.run(input_documents=docs, question=question)
    query_engine = vectorstore.as_query_engine(similariy_top_k=25)
    response = query_engine.query(question)
    return question, response


#CRM data handling


def parse_document(doc):
    data = {}
    current_section = None
    
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            if len(cells) == 2 and cells[0] == cells[1]:
                header_text = cells[0].lower()
                if "account information" in header_text:
                    current_section = "Account Information"
                    data[current_section] = {}
                elif "account contacts" in header_text:
                    current_section = "Account Contacts"
                    data[current_section] = []
                elif "related borrowings" in header_text:
                    current_section = "Related Borrowings"
                    data[current_section] = []
                continue
            
            if current_section and len(cells) >= 2:
                key, value = cells[0], cells[1]
                
                if isinstance(data[current_section], dict):
                    data[current_section][key] = value
                elif isinstance(data[current_section], list):
                    if key.lower() == "contact name":
                        data[current_section].append({key: value})
                    elif data[current_section]:
                        data[current_section][-1][key] = value
    
    return data


#Credit Application Handling


def update_section(application_id, section_name: str, new_data: dict):
    CREDIT_APPLICATION_STORE = {}
    if application_id not in CREDIT_APPLICATION_STORE:
        raise ValueError("Application ID not found")

    credit_application = CREDIT_APPLICATION_STORE[application_id]

    if not hasattr(credit_application, section_name):
        raise AttributeError(f"Section '{section_name}' not found in CreditApplication")

    section_obj = getattr(credit_application, section_name)

    if isinstance(section_obj, BaseModel):
        # Validate the new data before applying
        updated_section = section_obj.model_validate({**section_obj.model_dump(), **new_data})
        setattr(credit_application, section_name, updated_section)
    else:
        raise TypeError(f"Section '{section_name}' is not a Pydantic model")

    return credit_application

def create_credit_application(application_id):
    DATA_DIR = f"credit_applications/{application_id}.json"
    with open(DATA_DIR, "w") as f:
        json.dump(CreditApplication().model_dump(), f, indent=4)

def get_credit_application(application_id):
    try:
        with open(f"credit_applications\\{application_id}.json", "r") as f:
            data = json.load(f)
        return CreditApplication(**data)
    except FileNotFoundError:
        print("Credit Application not found for this application id")
        return None
    
def update_credit_application(application_id, new_credit_application_data):
    DATA_DIR = f"credit_applications/{application_id}.json"
    if os.path.exists(DATA_DIR):
        with open(DATA_DIR, "w") as f:
            json.dump(new_credit_application_data.model_dump(), f, indent=4)
    else:
        print(f"Could not find credit_application for {application_id}")

def update_credit_application_section(application_id, section_name, section_data):
    credit_application = get_credit_application(application_id)
    if credit_application is None:
        print(f"No CreditApplication found for ID {application_id}")
        return ValueError

    if not hasattr(credit_application, section_name):
        print(f"Section '{section_name}' not found in CreditApplication")
        return AttributeError
    
    section_value = getattr(credit_application, section_name)

    section_cls = type(section_value)

    if isinstance(section_value, BaseModel):
        updated_section = section_cls.model_validate(section_data)
    else:
        # For primitive types like str, int, etc., assign directly
        updated_section = section_data
    


    setattr(credit_application, section_name, updated_section)
    update_credit_application(application_id, credit_application)


#Word doc generation handling

def markdown_to_text(md: str) -> str:
    # Convert Markdown to HTML
    html = markdown.markdown(md)
    # Use BeautifulSoup to remove HTML tags and get plain text
    soup = BeautifulSoup(html, features="html.parser")
    return soup.get_text(separator="\n").strip()


def flatten_dict(d, parent_key=""):
    """Flattens a nested dictionary but keeps only the last-level key names."""
    flat_dict = {}

    def recurse(sub_d):
        if isinstance(sub_d, dict):
            for key, value in sub_d.items():
                if isinstance(value, dict):
                    recurse(value)  # Keep going deeper
                elif isinstance(value, list):
                    for i, item in enumerate(value):
                        if isinstance(item, dict):
                            recurse(item)  # Flatten each dict in list
                        else:
                            flat_dict[key] = value  # Keep non-dict lists as-is
                else:
                    flat_dict[key] = value  # Store final value

    recurse(d)
    return flat_dict

def generate_doc_response(application_id, DOCS_DIR = "docs", TEMPLATE_PATH = "assets\\placeholder_template.docx"):
    credit_application = get_credit_application(application_id)
    if not credit_application:
        raise ValueError("Invalid credit application object")

    output_path = os.path.join(DOCS_DIR, f"{application_id}_final.doc")

    # Load the DOCX template
    doc = Document(TEMPLATE_PATH)

    # Convert credit_application object to a dictionary & flatten it
    template_data = flatten_dict(credit_application.model_dump())

    # Replace placeholders in the DOCX
    for key, value in template_data.items():
        value = "" if not value else value
        replace_text_in_doc_para(doc, key, str(value))  # Ensure value is a string

    # Save the modified document
    doc.save(output_path)
    print(f"✅ Document saved at: {output_path}")
    return output_path

def replace_text_in_doc_para(doc, old_text, new_text):
    placeholder = f"##{old_text}##"

    new_text = markdown_to_text(new_text)

    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if placeholder not in paragraph.text:
                continue

            # Try replacing in individual runs first
            replaced = False
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    replaced = True

            # If not replaced in a single run, replace entire paragraph
            if not replaced:
                full_text = paragraph.text.replace(placeholder, new_text)
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                else:
                    paragraph.add_run(full_text)

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)
                    replace_in_tables(cell.tables)

    replace_in_paragraphs(doc.paragraphs)
    replace_in_tables(doc.tables)





#Unused


def generate_pdf_response(questionResponseMap, pdf_filename, application_id):
    # Create a PDF document
	#pdf_filename = f"responses/{application_id}_response.pdf"
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)

    # Create a style sheet for the document
    styles = getSampleStyleSheet()
    style_body = styles["BodyText"]
    style_heading = styles["Heading1"]
    style_subheading = styles["Heading3"]

    with open("constant\\constant.json", "r") as f:
        json_file = json.load(f)

    # Create paragraphs from response texts with proper text wrapping
    response_paragraphs = []
    for section in json_file["data"]:
        response_paragraphs.append(Paragraph(section["heading"] + "\n\n", style_heading))
        for question in section["questions"]:
            if isinstance(question, dict):
                label = list(question.keys())[0]
                response_paragraphs.append(Paragraph(label + " : ", style_subheading))
                response_paragraphs.append(Paragraph(questionResponseMap[question[label]], style_body))

            elif isinstance(question, str):
                response_paragraphs.append(Paragraph(question + " : ", style_subheading))
                response_paragraphs.append(Paragraph(questionResponseMap.get(question, "No response available"), style_body))

        response_paragraphs.append(Paragraph("\n\n", style_body))
 
    # Add paragraphs to the PDF document
    doc.build(response_paragraphs)
    #pdf_filename = f"pdfs/{application_id}_response.pdf"
    doc = fitz.open(pdf_filename)
    w = 595
    h = 842
    numpages = doc.page_count  # number of pages
    footer_text = "Page %i | Confidential"
    img = open("assets\\tcs_logo.png", "rb").read()
    logo_rect = fitz.Rect(w * 0.1, 0.03 * h, w * 0.3, h * 0.09)
    footer_rect = fitz.Rect(w * 0.4, h * 0.9, w * 0.65, h)

    for page in doc:
        if not page.is_wrapped:
            page.wrap_contents()
        page.insert_image(logo_rect, stream=img)
        f_text = footer_text % (page.number + 1)
        page.insert_textbox(footer_rect, f_text, align=fitz.TEXT_ALIGN_CENTER)
    #doc.save("pdfs//report.pdf")
    pdf_filename = f"pdfs/{application_id}_final.pdf"
    doc.save(pdf_filename)

def replace_text_in_doc(doc, old_text, new_text):
    placeholder = f"##{old_text}##"  # Ensure unique placeholders

    new_text = markdown_to_text(new_text)

    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    print("Replaced", run.text)

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)
                    replace_in_tables(cell.tables)  # Handle nested tables too

    replace_in_paragraphs(doc.paragraphs)
    replace_in_tables(doc.tables)